# Final project: Nólë 
#by Débora Machado Andrade, 27.03.2024


#(Nólë means "knowledge" in Quenya, a fictional Elvish language created by J.R.R.Tolkien)



import os
from openai import AsyncOpenAI  # importing openai for API usage
import chainlit as cl  # importing chainlit for our app
from chainlit.prompt import Prompt, PromptMessage  # importing prompt tools
from chainlit.playground.providers import ChatOpenAI  # importing ChatOpenAI tools
from story_generator import generate_story_from_info, generate_user_info_json, check_for_email, generate_summary_from_story
from rag_search_module import RAG_search
from image_prompt_module import process_text_for_image_prompts
from image_module2 import generate_image, download_image
import json
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches

document = Document()




load_dotenv()

# ChatOpenAI Templates
system_template = """You are a helpful children's book writer assistant. 
"""

user_template = """You are a helpful children's book writer assistant working for our publisher, which is specialized in educational books for children. 
Our books typically embed an educational or knowledge topic into an engaging narrative, in order to help children understand and assimilate these topics.
The user might be an educator or a parent, who has an idea for a book they wish to produce. 

Your job is to have a conversation with the user to understand what they are interested in writing. 
Your discussion with the user will help the next assistant to do research on the topic that will be explored in the book as well as the writer who will put all that together in an engaging story.
Regarding the educational topic embedded in the background, you should request references in the form of pdf documents or links. 
You will proactively ask questions, in a conversational fashion. 
You will only ask one question at a time. Your follow-up questions will depend on the user feedback. 
In the first part of your task, your aim is to end up with a clear idea of the following:

1) Summary of the plot
2) All main characters, their personalities, their physical appearance (if applicable), their personal story (if applicable)
3) Where and when the story happens
4) Which knowledge topic we cover in the background.
6) How can you choose an interesting ending.

Sometimes the user will not have concrete ideas regarding one or more aspects of the story and might ask for your help, in which case you should feel free to propose options that you find interesting and ask for feedback on that.
When and only when you gather all the information you need, (unless the user lets you know that they are leaving certain aspects of the story up to you), 
ask the user the following question exactly: 

Can you please provide a link to a pdf file that provides the educational content to be added to the story?

After you get a response, ask the following question exactly:

Can you please provide your email address?

Once your receive feedback about the user's email address, tell the user they will receive an email with a document containing their story in a few minutes. 
{input}. 
"""


@cl.on_chat_start  # marks a function that will be executed at the start of a user session

async def start_chat():
    settings = {
        "model": "gpt-4",
        "temperature": 1.0,
        "max_tokens": 900,
        "top_p": 1,
        "frequency_penalty": 0,
        "presence_penalty": 0,
    }

    cl.user_session.set("settings", settings)
    cl.user_session.set("conversation_history", [])  # Initialize empty conversation history

@cl.on_message
async def main(message: cl.Message):
    settings = cl.user_session.get("settings")
    conversation_history = cl.user_session.get("conversation_history")

    client = AsyncOpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

    # Append the new user message to the conversation history
    conversation_history.append({"role": "user", "content": message.content})

    # Constructing a dynamic prompt using the conversation history
    dynamic_messages = []
    for msg in conversation_history:
        if msg["role"] == "system":
            dynamic_messages.append(PromptMessage(role="system", template=system_template, formatted=system_template))
        else:
            dynamic_messages.append(PromptMessage(role="user", template=user_template, formatted=user_template.format(input=msg["content"])))

    dynamic_prompt = Prompt(
        provider=ChatOpenAI.id,
        messages=dynamic_messages,
        inputs={"input": message.content},
        settings=settings,
    )

    #print("Sending to OpenAI:", [m.to_openai() for m in dynamic_prompt.messages])

    msg = cl.Message(content="")



    # Call OpenAI with the dynamic prompt
    async for stream_resp in await client.chat.completions.create(
        messages=[m.to_openai() for m in dynamic_prompt.messages], stream=True, **settings
    ):
        token = stream_resp.choices[0].delta.content
        if not token:
            token = ""
        await msg.stream_token(token)

    # Update the prompt object with the completion and add it to the conversation history
    dynamic_prompt.completion = msg.content
    conversation_history.append({"role": "system", "content": msg.content})
    
    # Save the updated conversation history back to the session
    cl.user_session.set("conversation_history", conversation_history)


    if check_for_email(message.content):
    #triggers the generation of the story and images
        
        #First, the conversation is summarized into a json:
        user_info_json = await generate_user_info_json(client, conversation_history) 
        #print(user_info_json)

        #The pdf file the user referred to is downloaded and used to enrich the information on the educational topic:
        user_info = json.loads(user_info_json)
        pdf = user_info["pdf"]
        search_topic = user_info["search_topic"]
        characters_appearance = user_info["characters_physical_appearance"]

        information = await RAG_search(pdf, search_topic)

        #The story is generated based on the user requests and the information retrieved from the pdf, then saved in a docx file:
        story = await generate_story_from_info(client, user_info, information)
        summary = await generate_summary_from_story(client, story)
        print(summary)
        #The story is split into paragraphs and those are processed for image prompting:
        rewritten_paragraphs, paragraphs = await process_text_for_image_prompts(client, story)

        #For each image prompt (which means for each paragraph), we generate an image and save it in the local directory.
        #In addition, we add each paragraph at a time followed by its corresponding image to our docx file:
         
        for i in range(len(paragraphs)):
            filename = await generate_image(client, rewritten_paragraphs[i], characters_appearance)
            document.add_paragraph(paragraphs[i])
            document.add_picture(filename, width=Inches(4))
            document.save('story.docx')
        

    # Send and close the message stream
    await msg.send()

